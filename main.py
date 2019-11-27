from __future__ import annotations

import argparse
import functools
import re
import unicodedata
import logging
from enum import Enum
from pathlib import Path
from re import search
from typing import List, Optional, Dict, Tuple

from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTPage, LTTextBoxHorizontal, LTTextLineHorizontal
from pdfminer.pdfdevice import PDFDevice
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFPageInterpreter, PDFResourceManager
from pdfminer.pdfpage import PDFPage, PDFTextExtractionNotAllowed
from pdfminer.pdfparser import PDFParser
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def uni_nornal(func):
    """ 替换`\xa0`为space """

    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        res = func(*args, **kwargs)
        return unicodedata.normalize('NFKD', res)

    return wrapper


class SubjectType(Enum):
    cloze = '完形填空'
    read = '阅读理解'
    read_75 = '阅读理解7选5'
    fill = '语法填空'
    translate = '论述题'


class PDFLine:
    debug = False
    _types = [type.value for type in SubjectType]

    def __init__(self, index: str, line: LTTextLineHorizontal):
        self._obj = line
        self.index = index
        self.text = self.get_text()

    def __repr__(self):
        return f'{self.index}\n{self.text}'

    @property
    def is_type(self):
        """ 是否内容是类型字段 """
        return self.text.startswith(tuple(self._types))

    @property
    def is_new_paragraph(self):
        """ 是否新段落开始 """
        # 1. 自然段起始四(三)空格+大写字母(兼容3空格形式)
        # 2. 子标题: 1. Xxxxx
        return re.search(r'^\s{3,4}[A-Z]', self.text) or re.search(
            r'^\d\.\s[A-Z]', self.text
        )

    @property
    def is_other_line(self):
        """ 是否其余段 """
        # 不能是`x.`形式
        return not self.is_new_paragraph

    @property
    def is_index(self):
        """ 是否完型的题序号 """
        # `x. $`
        return re.search(r'^\d+\.$', self.text)

    @property
    def is_index_with_qst(self):
        """ 是否带题序号 """
        # `x. xxxxxxx`
        return re.search(r'^\d+\.\s.+', self.text)

    @property
    def is_option(self):
        """ 是否选项 """
        # `A. xxxxxxxxxxxxx`
        return re.search(r'^[A-G]\..+', self.text)

    def get_text(self):
        res = self._obj.get_text()
        # 替换`\xa0`为` `, 去除末尾空格和换行(\n)
        res = unicodedata.normalize('NFKD', res).rstrip()
        return res


class PDFBox:
    debug = False

    def __init__(self, box: LTTextBoxHorizontal, page: int, last_box: Optional[PDFBox]):
        self._obj = box
        self._last = last_box
        self.page = page
        self.index: str = f'{page}-{box.index}'
        # self.text = box.get_text()
        self.lines = self.get_lines()
        self.size = len(self.lines)
        self.pars = self.convert_lines_to_pars()

    def __repr__(self):
        return f'--- {self.index} --- {type(self._obj).__name__} ---\n'

    @property
    def is_multilines(self):
        """ 是否内容多行 """
        return len(self.lines) > 1

    @property
    def is_new_subject(self):
        """ 是否新题开始 """
        # 1. 首行(上行不存在)
        # ABANDONED: 2. 上行 `x. xxxxx`格式 & 当前行不是此格式
        # 2. 两篇文章分界
        return self._last is None or (not self._last.is_body and self.is_body)

    @property
    def is_body(self):
        # 包含4个空格
        # `\s?`防止`I xxxx`类型
        return re.search(r'\s{3,4}[A-Z]\s?\w+', self.text)

    @property
    @uni_nornal
    def text(self):
        if self.debug is True:
            return self._obj.get_text()
        return ''.join(self.pars)

    def get_lines(self):
        """ 获取box的所有行对象 """
        _lines: List[PDFLine] = []
        for i, line in enumerate(self._obj):
            _lines.append(PDFLine(f'{self.index}-{i}', line))
        return _lines

    def convert_lines_to_pars(self):
        """ 遍历行对象生成段落列表 """
        _pars: List[str] = []
        for line in self.lines:
            # if line.is_first_line:
            # _pars.append(line.text)
            # 当前是` xxx`开头 或上行末尾是`.$`
            if len(_pars) > 0 and (
                re.search(r'^\s\S', line.text) or not _pars[-1][-2:] == '.\n'
            ):
                # 加入当前最后一段
                _pars[-1] = _pars[-1][:-1] + line.text
            else:
                _pars.append(line.text)
        return _pars


class Subject:
    # TODO: 更改为enm选项
    type: Optional[str]
    titles: List[str]
    paragraphs: List[str]
    qsts: List[str]
    options: {}

    def __init__(self):
        self._body = []
        self._foots = []

        self.type = None
        self.titles = []
        self.paragraphs = []
        self.qsts = []
        self.indexs = []
        self.options = {
            'A': [],
            'B': [],
            'C': [],
            'D': [],
        }

        self.is_has_p = False

    def is_title(self, line: PDFLine):
        return not self.is_has_p and line.is_other_line

    def is_next_line(self, line: PDFLine):
        return self.is_has_p and line.is_other_line

    def classify_box(self, box: PDFBox):
        """ 分类PDFBox对象, 加入队列"""
        # if not self.is_has_head:
        # logging.info(f'Sub-parse-head: {box.text}')
        # self._parse_head.append(box)
        if box.is_body:
            logging.info(f'Sub-parse-body: {repr(box)}\n{box.text}')
            self.is_has_body = True
            self._body.append(box)
        else:
            logging.info(f'Sub-parse-foot: {repr(box)}\n{box.text}')
            self._foots.append(box)

    def parse_all(self):
        self.parse_body(self._body)
        self.parse_foots(self._foots)
        return True

    def parse_body(self, boxes: List[PDFBox]):
        for box in boxes:
            for line in box.lines:
                if line.is_type:
                    logging.debug(f'type: {line}')
                    self.type = line.text
                elif self.is_title(line):
                    logging.debug(f'title: {line}')
                    self.titles.append(line.text)
                elif line.is_new_paragraph:
                    logging.debug(f'para: {line}')
                    if self.is_has_p is False:
                        self.is_has_p = True
                        logging.debug('para is first para')
                    # 加入, 去除头部空格
                    self.paragraphs.append(line.text.strip())
                elif self.is_next_line(line):
                    logging.debug(f'next line: {line}')
                    self.paragraphs[-1] += line.text
                else:
                    logging.warning(f'NOT MATCH: {line}\t{line.text}')

    def parse_foots(self, boxes: List[PDFBox]):
        """ 解析问题和选项 """
        _lines = []
        for box in boxes:
            _lines += box.lines

        for line in _lines:
            if line.is_index:
                logging.debug(f'qst without index: {line}')
                self.indexs.append(line.text)
            elif line.is_index_with_qst:
                logging.debug(f'qst with index: {line}')
                self.qsts.append(line.text)
            elif line.is_option:
                # 如果是完型
                if self.type == SubjectType.cloze.value:
                    logging.debug(f'cloze option: {line}')
                    char = line.text[0]
                    self.options[char].append(line.text)
                # 如果阅读
                elif self.type == SubjectType.read.value:
                    logging.debug(f'read option: {line}')
                    # 如果长度过长, 可能存在2个选项识别为一行的情况
                    if len(line.text) > 50:
                        logging.debug(f'option too long: {len(line.text)}')
                        # `A. xxx B. xxx`
                        # to
                        # `A. xxxx`
                        # `B. xxxx`
                        t = re.sub(r'(?<!^)\s+(?=[A-D]\.\s)', '`', line.text)
                        tlist = t.split('`')
                        logging.debug(f'split to: {t}')

                        for t in tlist:
                            char = t[0]
                            self.options[char].append(t)
                    else:
                        char = line.text[0]
                        self.options[char].append(line.text)
                # 如果是7选5
                elif self.type == SubjectType.read_75.value:
                    logging.debug(f'7to5 option: {line}')
                    self.qsts.append(line.text)
            else:
                logging.warning(f'NOT MATCH TYPE: {self.type}')

    def check_and_fix_parse_seq(self, parameter_list):
        pass

    def add(self, box: PDFBox):
        for line in box.lines:
            if line.is_type:
                logging.debug(f'type: {line}')
                self.type = line.text
            elif self.is_title(line):
                logging.debug(f'title: {line}')
                self.titles.append(line.text)
            elif line.is_new_paragraph:
                if not self.is_has_p:
                    self.is_has_p = True
                logging.debug(f'para: {line}')
                # 加入, 去除头部空格
                self.paragraphs.append(line.text.strip())
            elif line.is_index:
                logging.debug(f'qst without index: {line}')
                self.indexs.append(line.text)
            elif line.is_index_with_qst:
                # 排除7选5正文中1. xxxx的内容
                if self.type != SubjectType.read_75.value:
                    logging.debug(f'next line(7to5_fix): {line}')
                    self.paragraphs[-1] += line.text
                else:
                    logging.debug(f'qst with index: {line}')
                    self.qsts.append(line.text)
            elif self.is_next_line(line):
                logging.debug(f'next line: {line}')
                self.paragraphs[-1] += line.text
            elif line.is_option:
                # 如果是完型
                if self.type == SubjectType.cloze.value:
                    logging.debug(f'cloze option: {line}')
                    char = line.text[0]
                    self.options[char].append(line.text)
                # 如果阅读
                elif self.type == SubjectType.read.value:
                    logging.debug(f'read option: {line}')
                    # 如果长度过长, 可能存在2个选项识别为一行的情况
                    if len(line.text) > 50:
                        logging.debug(f'option too long: {len(line.text)}')
                        # `A. xxx B. xxx`
                        # to
                        # `A. xxxx`
                        # `B. xxxx`
                        t = re.sub(r'(?<!^)\s+(?=[A-D]\.\s)', '`', line.text)
                        tlist = t.split('`')
                        logging.debug(f'split to: {t}')

                        for t in tlist:
                            char = t[0]
                            self.options[char].append(t)
                    else:
                        char = line.text[0]
                        self.options[char].append(line.text)
                # 如果是7选5
                elif self.type == SubjectType.read_75.value:
                    logging.debug(f'7to5 option: {line}')
                    self.qsts.append(line.text)
                else:
                    logging.warning(f'NOT MATCH TYPE: {self.type}')

            else:
                logging.warning(f'NOT MATCH: {line.text}')

    def get_title(self):
        """ 获取标题 """
        return '\n'.join(self.titles)

    def get_p_text(self):
        """ 获取段落文本 """
        res = ''
        for p in self.paragraphs:
            logging.debug(f'Get p text: {repr(p)}')
            if self.type == SubjectType.fill.value:
                # 语法填空格式替换为`x_____`

                # 匹配` xx(` 或 ` xx `
                p = re.sub(r'\s+(\d+)\s+', r' \1_________ ', p)
                p = re.sub(r'\s+(\d+)(?=\()', r' \1_________ ', p)

                p = re.sub(r'\s*(\d+)\s{2,}', r' \1_________ ', p)
            else:
                # 修复文章中题序号为_x_格式
                # re: 前后为字符+若干空格
                # TODO: 没有匹配`" 30 "`形式
                p = re.sub(r'\s{2,}(\d+)\s*', r' _\1_ ', p)
                p = re.sub(r'\s*(\d+)\s{2,}', r' _\1_ ', p)

            res += ' ' * 4 + p + '\n'

        logging.debug(f'return content text: {res}')
        return res

    def get_qst_text(self):
        res = ''
        align = 25
        if self.type == SubjectType.cloze.value:
            idx = self.indexs
            a = self.options['A']
            b = self.options['B']
            c = self.options['C']
            d = self.options['D']
            for i in range(len(idx)):
                qostr = f"{idx[i]:<5}{a[i]:<25}{b[i]:<25}{c[i]:<25}{d[i]}\n"
                logging.debug(qostr)
                res += qostr
        if self.type == SubjectType.read.value:
            q = self.qsts
            a = self.options['A']
            b = self.options['B']
            c = self.options['C']
            d = self.options['D']
            break_char = ''
            for i in range(len(q)):
                qstr = f'{q[i]}'
                logging.debug(f'qestion:\n{qstr}')
                if len(a[i] + b[i] + c[i] + d[i]) > 100:
                    align = 56
                    # 分两行显示
                    break_char = '\n'
                opt_str = (
                    f"{a[i]:<{align}}{b[i]:<{align}}{break_char}{c[i]:<{align}}{d[i]}"
                )
                logging.debug(f'option:\n{opt_str}')
                res += qstr + '\n' + opt_str + '\n'
        if self.type == SubjectType.read_75.value:
            # qsts里是重复的同7个选项, 第一项是序号, 取1-7项即为A-G
            res = '\n' + '\n'.join(sorted(self.qsts[1:8])) + '\n'
        return res

    def get_text(self):
        if self.type == SubjectType.translate.value:
            _type = '短文改错'
        else:
            _type = self.type
        return (
            _type
            + '\n'
            + self.get_title()
            + '\n'
            + self.get_p_text()
            + self.get_qst_text()
        )


def parse_file(file: Path):
    with open(file, 'rb') as fp:
        parser = PDFParser(fp)
        doc = PDFDocument(parser)
        laparams = LAParams()
        text_boxes = []  # 清理后box列表

        if not doc.is_extractable:
            raise PDFTextExtractionNotAllowed

        rsrcmgr = PDFResourceManager()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        last_out = None
        for i, page in enumerate(PDFPage.create_pages(doc)):
            orgi_boxes = []  # 原始box列表

            interpreter.process_page(page)
            layout = device.get_result()

            for out in layout:
                if isinstance(out, LTTextBoxHorizontal) and (
                    '猿题库' not in out.get_text()
                ):
                    orgi_boxes.append(out)
                else:
                    pass
            # 去除页眉和页尾
            cleaned_boxes = orgi_boxes[1:-1]

            for box in cleaned_boxes:
                pdf_box = PDFBox(box, i, last_out)
                text_boxes.append(pdf_box)
                last_out = pdf_box
    print('parse end')
    return text_boxes


def concat_subjects(boxes: List[PDFBox]):
    res: List[Subject] = []
    sub = None

    for box in boxes:
        if box.is_new_subject:
            sub = Subject()
            res.append(sub)
            logging.info(f'Create new subject:\n{repr(box.text)}')
            # sub.add(box)
            # if sub.type:
            # last_type = sub.type
            # logging.debug(f'Use New Type: {sub.type}')
            # else:
            # sub.type = last_type
            # logging.debug(f'Use Last Type: {sub.type}')
        # else:
        # sub.add(box)
        # logging.info(f'Classify box:\n{repr(box.text)}')
        sub.classify_box(box)

    for i, s in enumerate(res):
        # 默认后续文章使用上篇文章的类型
        if i > 0:
            s.type = res[i - 1].type
        s.parse_all()

    return res


def write_file(objs: List[Subject], src_file: Path, filename: Path = None, flags='w'):
    result = False
    if filename is None:
        filename = src_file.with_suffix('.txt')

    print(filename)

    try:
        with open(filename, flags) as fp:
            for obj in objs:
                fp.writelines(obj.get_text())
            result = True
    except IOError as identifier:
        print(identifier)
    return result


def write_to_docx(
    objs: List[Subject], src_file: Path, filename: Path = None, flags='w'
):
    doc = Document()
    set_doc_format(doc)

    for obj in objs:
        doc.add_paragraph(obj.type, 'TYPE')
        for title in obj.titles:
            doc.add_paragraph(title, 'TITLE')
        if obj.type == SubjectType.translate.value:
            doc.add_paragraph(obj.get_p_text(), 'TRANSLATE')
        else:
            doc.add_paragraph(obj.get_p_text())
        doc.add_paragraph(obj.get_qst_text())

    if filename is None:
        filename = src_file.with_suffix('.docx')
    doc.save(filename)


def set_doc_format(doc: Document):
    # page
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

    # type
    type_style = doc.styles.add_style('TYPE', WD_STYLE_TYPE.PARAGRAPH)
    type_pf = type_style.paragraph_format
    type_pf.space_before = Pt(0)
    type_pf.space_after = Pt(0)
    type_style.font.size = Pt(14)
    # title
    title_style = doc.styles.add_style('TITLE', WD_STYLE_TYPE.PARAGRAPH)
    title_pf = title_style.paragraph_format
    title_pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_pf.space_before = Pt(0)
    title_pf.space_after = Pt(0)
    title_style.font.size = Pt(14)

    # normal
    normal_style = doc.styles['Normal']
    font = normal_style.font
    font.size = Pt(12)

    # translate
    tr_style = doc.styles.add_style('TRANSLATE', WD_STYLE_TYPE.PARAGRAPH)
    tr_style.paragraph_format.line_spacing = Pt(24)
    font = tr_style.font
    font.size = Pt(12)

    return doc


def main():
    # encoding = 'utf-8'

    FORMAT = (
        '%(asctime)s,%(msecs)d %(levelname)-8s [%(filename)s:%(lineno)d] %(message)s'
    )
    logger = logging.getLogger()
    handler = logging.FileHandler('debug.log', 'w')
    formatter = logging.Formatter(FORMAT, datefmt='%Y-%m-%d:%H:%M:%S')
    handler.setFormatter(formatter)
    logger.addHandler(handler)

    parser = argparse.ArgumentParser()
    parser.add_argument('file', nargs='?', default='t.pdf')
    parser.add_argument('-d', dest='debug', action='store_true', help='debug on')

    args = parser.parse_args()

    print(args)
    PDFBox.debug = args.debug
    PDFLine.debug = args.debug
    if args.debug:
        logger.setLevel(logging.DEBUG)

    logging.info('\n-----------------Start---------------\n')

    _file = Path(args.file)
    res_text = parse_file(_file)
    subjects = concat_subjects(res_text)

    write_file(subjects, _file, Path('output.txt'))
    write_to_docx(subjects, _file)
    print(_file)


if __name__ == '__main__':
    main()
